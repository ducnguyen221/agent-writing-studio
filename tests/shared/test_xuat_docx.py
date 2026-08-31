"""`shared/scripts/xuat_docx.py` — bản giao ra Word có đúng quy cách không?

Hai tầng, tách rời có chủ đích:

1. **Parser** — thuần Python, chạy được cả trên máy chưa cài `python-docx`. Đây là chỗ dễ hỏng im
   lặng nhất: nuốt frontmatter nhầm, ăn mất một sao của `***`, hoặc **diễn giải chữ trong khối
   code** (code là vùng bảo vệ của trục 4 — sửa khéo trong đó là sửa sai).
2. **Quy cách** — Times New Roman 13pt, giãn dòng 1,5, lề 2/2/3/2 cm, heading đậm 14–16. Kiểm bằng
   cách dựng thật một `Document` rồi đọc lại thuộc tính, không tin vào hằng số trong nguồn: hằng số
   đúng mà quên áp vào section thì file giao ra vẫn sai lề.

Tầng 2 `skipUnless` khi thiếu `python-docx` — thư viện là **tuỳ chọn** theo README, nên test đỏ vì
môi trường thiếu thư viện là test nói dối về code.
"""

import importlib.util
import tempfile
import unittest
from pathlib import Path


ROOT = Path(__file__).resolve().parents[2]
SCRIPT = ROOT / "shared/scripts/xuat_docx.py"


def load_module():
    spec = importlib.util.spec_from_file_location("xuat_docx", SCRIPT)
    module = importlib.util.module_from_spec(spec)
    assert spec.loader is not None
    spec.loader.exec_module(module)
    return module


XD = load_module()

try:  # thư viện tuỳ chọn — xem docstring
    import docx as _docx  # noqa: F401

    CO_DOCX = True
except ImportError:  # pragma: no cover - phụ thuộc môi trường
    CO_DOCX = False


MAU = """---
title: không phải nội dung
---

# Tiêu đề

Đoạn một
nối dòng mềm.

- gạch đầu dòng
1. đánh số

> trích dẫn

```python
x = 1  # ** đây không phải chữ đậm **
```
"""


class ParserTests(unittest.TestCase):
    def test_bo_frontmatter_yaml(self):
        loai = [khoi["loai"] for khoi in XD.phan_tich_markdown(MAU)]
        self.assertEqual(loai[0], "heading")
        chu = " ".join(
            doan.chu
            for khoi in XD.phan_tich_markdown(MAU)
            for doan in khoi.get("doan", [])
        )
        self.assertNotIn("không phải nội dung", chu)

    def test_nhan_dung_sau_loai_khoi(self):
        loai = [khoi["loai"] for khoi in XD.phan_tich_markdown(MAU)]
        self.assertEqual(
            loai,
            ["heading", "doan", "gach_dau_dong", "danh_so", "trich_dan", "code"],
        )

    def test_dong_mem_gop_thanh_mot_doan(self):
        doan = [k for k in XD.phan_tich_markdown(MAU) if k["loai"] == "doan"][0]
        self.assertEqual(doan["doan"][0].chu, "Đoạn một nối dòng mềm.")

    def test_heading_sau_hon_ba_muc_quy_ve_muc_ba(self):
        khoi = XD.phan_tich_markdown("##### rất sâu\n")
        self.assertEqual(khoi[0]["muc"], 3)

    def test_khoi_code_giu_nguyen_van_khong_dien_giai(self):
        code = [k for k in XD.phan_tich_markdown(MAU) if k["loai"] == "code"][0]
        self.assertIn("** đây không phải chữ đậm **", code["chu"])

    def test_dam_va_nghieng_long_nhau(self):
        doan = XD.chia_inline("Câu **đậm và *nghiêng*** rồi.")
        self.assertIn(("đậm và ", True, False, False), [d.as_tuple() for d in doan])
        self.assertIn(("nghiêng", True, True, False), [d.as_tuple() for d in doan])

    def test_ma_inline_va_lien_ket(self):
        doan = XD.chia_inline("dùng `lệnh` xem [trang](https://vi.dụ) nhé")
        bo = [d.as_tuple() for d in doan]
        self.assertIn(("lệnh", False, False, True), bo)
        self.assertIn(("trang", False, False, False), bo)
        self.assertIn((" (https://vi.dụ)", False, False, False), bo)

    def test_lien_ket_noi_bo_khong_bi_dan_duong_dan_vao_van(self):
        """URL ngoài thì in ra cho người đọc tra được; đường dẫn file nội bộ thì vô nghĩa với họ."""
        doan = XD.chia_inline("xem [tài liệu](docs/KIEN-TRUC.md)")
        self.assertNotIn("docs/KIEN-TRUC.md", " ".join(d.chu for d in doan))


class DuongDanDichTests(unittest.TestCase):
    def test_out_la_thu_muc_thi_lay_ten_theo_nguon(self):
        self.assertEqual(
            XD.duong_dan_dich("ca/polished.md", "C:/nguoi-dung/bai"),
            Path("C:/nguoi-dung/bai/polished.docx"),
        )

    def test_out_la_file_docx_thi_giu_nguyen(self):
        self.assertEqual(
            XD.duong_dan_dich("ca/polished.md", "C:/nguoi-dung/ban-giao.docx"),
            Path("C:/nguoi-dung/ban-giao.docx"),
        )

    def test_out_co_duoi_khac_thi_bao_loi_thay_vi_doan(self):
        with self.assertRaises(ValueError):
            XD.duong_dan_dich("ca/polished.md", "C:/nguoi-dung/ban-giao.pdf")


class ThieuThuVienTests(unittest.TestCase):
    def test_thong_bao_thieu_thu_vien_chi_ro_cach_cai(self):
        self.assertIn("pip install python-docx", XD.THIEU_THU_VIEN)


@unittest.skipUnless(CO_DOCX, "cần python-docx (thư viện tuỳ chọn)")
class QuyCachTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        cls.tai_lieu = XD.dung_tai_lieu(XD.phan_tich_markdown(MAU))

    def test_le_dung_chuan_van_ban_viet(self):
        phan = self.tai_lieu.sections[0]
        self.assertAlmostEqual(phan.top_margin.cm, 2.0, places=2)
        self.assertAlmostEqual(phan.bottom_margin.cm, 2.0, places=2)
        self.assertAlmostEqual(phan.left_margin.cm, 3.0, places=2)
        self.assertAlmostEqual(phan.right_margin.cm, 2.0, places=2)

    def test_font_co_chu_va_gian_dong_mac_dinh(self):
        binh_thuong = self.tai_lieu.styles["Normal"]
        self.assertEqual(binh_thuong.font.name, "Times New Roman")
        self.assertEqual(binh_thuong.font.size.pt, 13)
        self.assertEqual(binh_thuong.paragraph_format.line_spacing, 1.5)

    def test_moi_doan_deu_duoc_ap_gian_dong(self):
        """Style Normal đúng chưa đủ: `List Bullet` mang paragraph_format riêng và sẽ đè lên."""
        for doan_van in self.tai_lieu.paragraphs:
            if doan_van.runs and doan_van.runs[0].font.name != "Consolas":
                self.assertEqual(doan_van.paragraph_format.line_spacing, 1.5)

    def test_heading_dam_va_co_chu_trong_khoang_14_16(self):
        heading = self.tai_lieu.paragraphs[0].runs[0]
        self.assertTrue(heading.bold)
        self.assertEqual(heading.font.size.pt, 16)
        for co_chu in XD.QUY_CACH["co_heading"].values():
            self.assertGreaterEqual(co_chu, 14)
            self.assertLessEqual(co_chu, 16)

    def test_danh_sach_dung_style_cua_word(self):
        style = [p.style.name for p in self.tai_lieu.paragraphs]
        self.assertIn("List Bullet", style)
        self.assertIn("List Number", style)


@unittest.skipUnless(CO_DOCX, "cần python-docx (thư viện tuỳ chọn)")
class ChayThatTests(unittest.TestCase):
    def test_xuat_ra_file_that_va_chep_sidecar_provenance(self):
        with tempfile.TemporaryDirectory() as tam:
            tam = Path(tam)
            nguon = tam / "polished.md"
            nguon.write_text(MAU, encoding="utf-8")
            sidecar = tam / "polished.provenance.json"
            sidecar.write_text('{"stylometric_polish": true}\n', encoding="utf-8")

            dich = XD.xuat(nguon, tam / "ban-giao", provenance=sidecar)

            self.assertTrue(dich.is_file(), "không ghi được file docx")
            self.assertEqual(dich.name, "polished.docx")
            self.assertGreater(dich.stat().st_size, 0)
            self.assertTrue(
                (dich.parent / "polished.provenance.json").is_file(),
                "sidecar provenance phải nằm CẠNH bản giao, không ở lại thư mục ca",
            )


if __name__ == "__main__":
    unittest.main()
