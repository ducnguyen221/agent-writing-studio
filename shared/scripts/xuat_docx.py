#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
xuat_docx.py — Markdown ➜ .docx theo quy cách văn bản tiếng Việt.

VÌ SAO CÓ FILE NÀY. Studio làm việc bằng Markdown, nhưng **bản giao cho người đọc mặc định là
`.docx`**: giảng viên, biên tập viên và người quản lý nhận bài bằng Word, không bằng `.md`. Harness
trên máy đọc được docx (markitdown/pandoc) nhưng **không tạo được docx đúng quy cách Việt** — nên
phần tạo nằm ở đây, cạnh các script khác của studio.

Quy cách mặc định (`QUY_CACH`): Times New Roman 13pt · giãn dòng 1,5 · lề trên/dưới 2cm, trái 3cm,
phải 2cm · heading đậm cỡ 16/15/14. Đây là quy cách văn bản hành chính - học thuật phổ thông ở Việt
Nam; đổi nó là đổi một hằng số ở một chỗ, không phải sửa rải rác.

Cách chạy:

    python xuat_docx.py bai.md --out "C:/thu-muc-cua-nguoi-dung"
    python xuat_docx.py bai.md --out "C:/thu-muc/ten-ban-giao.docx" --provenance ca/polished.provenance.json

`--out` nhận **thư mục** (đặt tên file theo tên file nguồn) hoặc **đường dẫn `.docx`** cụ thể.
`--provenance` chép sidecar tự khai nguồn gốc sang **cạnh** file docx, đúng luật "provenance đi theo
bản giao" (`docs/KIEN-TRUC.md` §2.5 mục 5). Không có sidecar thì bản giao không tự chứng minh được
nguồn gốc — script chỉ cảnh báo, không tự bịa ra một file tự khai.

PHẦN MARKDOWN ĐƯỢC HIỂU (parser tự viết, không thêm phụ thuộc): frontmatter YAML (bỏ qua), heading
`#`…`###` (sâu hơn quy về mức 3), đoạn văn, danh sách gạch đầu dòng và danh sách đánh số, trích dẫn
`>`, khối code ```…``` (giữ nguyên văn, KHÔNG diễn giải), và inline `**đậm**`, `*nghiêng*`, `` `mã` ``,
`[chữ](đường-dẫn)`.

GIỚI HẠN ĐÃ BIẾT, ghi ra để không ai đọc nhầm: **bảng Markdown không được dựng thành bảng Word** —
dòng `| … |` đi ra ở dạng đoạn nguyên văn, đọc được nhưng không đẹp. Bài giao có bảng thì dùng
pandoc, hoặc dựng bảng trong Word sau. Ảnh, footnote và HTML thô cũng không được dịch.

Mã thoát: 0 = xong · 1 = lỗi (thiếu `python-docx`, không thấy file nguồn, `--out` sai kiểu).
"""
import argparse
import re
import shutil
import sys
from pathlib import Path

EXIT_OK, EXIT_LOI = 0, 1

QUY_CACH = {
    "font": "Times New Roman",
    "co_chu": 13,
    "gian_dong": 1.5,
    "le_tren_cm": 2.0,
    "le_duoi_cm": 2.0,
    "le_trai_cm": 3.0,
    "le_phai_cm": 2.0,
    # Heading đậm, cỡ nằm trong khoảng 14–16 — mức 1 to nhất.
    "co_heading": {1: 16, 2: 15, 3: 14},
    "font_ma": "Consolas",
}

# ---------------------------------------------------------------- parser

HEADING = re.compile(r"^(#{1,6})\s+(.*)$")
GACH_DAU_DONG = re.compile(r"^\s*[-*+]\s+(.*)$")
DANH_SO = re.compile(r"^\s*\d+[.)]\s+(.*)$")
TRICH_DAN = re.compile(r"^\s*>\s?(.*)$")
DUONG_KE = re.compile(r"^\s*(-{3,}|_{3,}|\*{3,})\s*$")
RAO_CODE = re.compile(r"^\s*(```|~~~)")

INLINE = re.compile(
    r"(?P<ma>`[^`\n]+`)"
    # `(?!\*)` để `**đậm và *nghiêng***` không dừng ở hai sao đầu của cụm ba sao rồi bỏ rơi một sao.
    r"|(?P<dam>\*\*(?P<dam_trong>.+?)\*\*(?!\*)|__(?P<dam_trong2>.+?)__(?!_))"
    r"|(?P<nghieng>\*(?P<nghieng_trong>[^*\n]+?)\*|_(?P<nghieng_trong2>[^_\n]+?)_)"
    r"|(?P<lien_ket>\[(?P<lk_chu>[^\]]*)\]\((?P<lk_dich>[^)\s]*)[^)]*\))",
)


class Doan:
    """Một mẩu chữ liền mạch cùng kiểu — đơn vị nhỏ nhất mà docx gọi là `run`."""

    __slots__ = ("chu", "dam", "nghieng", "ma")

    def __init__(self, chu, dam=False, nghieng=False, ma=False):
        self.chu = chu
        self.dam = dam
        self.nghieng = nghieng
        self.ma = ma

    def as_tuple(self):
        return (self.chu, self.dam, self.nghieng, self.ma)

    def __eq__(self, other):
        if isinstance(other, Doan):
            return self.as_tuple() == other.as_tuple()
        if isinstance(other, tuple):
            return self.as_tuple() == other
        return NotImplemented

    def __repr__(self):
        return f"Doan{self.as_tuple()!r}"


def chia_inline(chu, dam=False, nghieng=False, ma=False, sau=0):
    """Tách một dòng thành các `Doan` theo dấu đậm/nghiêng/mã/liên kết.

    Đệ quy để `**câu *nhấn* trong đậm**` giữ được cả hai kiểu; `sau` chặn đệ quy vô hạn nếu
    một mẫu bệnh hoạn nào đó khớp rỗng.
    """
    if not chu:
        return []
    if sau > 4:
        return [Doan(chu, dam, nghieng, ma)]

    ket_qua = []
    vi_tri = 0
    for khop in INLINE.finditer(chu):
        if khop.start() > vi_tri:
            ket_qua.append(Doan(chu[vi_tri:khop.start()], dam, nghieng, ma))
        if khop.group("ma"):
            ket_qua.append(Doan(khop.group("ma").strip("`"), dam, nghieng, True))
        elif khop.group("dam"):
            trong = khop.group("dam_trong") or khop.group("dam_trong2") or ""
            ket_qua += chia_inline(trong, True, nghieng, ma, sau + 1)
        elif khop.group("nghieng"):
            trong = khop.group("nghieng_trong") or khop.group("nghieng_trong2") or ""
            ket_qua += chia_inline(trong, dam, True, ma, sau + 1)
        else:
            ket_qua += chia_inline(khop.group("lk_chu") or "", dam, nghieng, ma, sau + 1)
            dich = khop.group("lk_dich") or ""
            if dich.startswith(("http://", "https://", "mailto:")):
                ket_qua.append(Doan(f" ({dich})", dam, nghieng, ma))
        vi_tri = khop.end()
    if vi_tri < len(chu):
        ket_qua.append(Doan(chu[vi_tri:], dam, nghieng, ma))
    return [doan for doan in ket_qua if doan.chu]


def _bo_frontmatter(dong_list):
    """Frontmatter YAML là metadata cho agent, không phải nội dung bản giao."""
    if dong_list and dong_list[0].strip() == "---":
        for so, dong in enumerate(dong_list[1:], start=1):
            if dong.strip() in ("---", "..."):
                return dong_list[so + 1:]
    return dong_list


def phan_tich_markdown(van_ban):
    """Markdown ➜ danh sách khối `{"loai": …, "muc": …, "doan": [...]}`.

    `loai` ∈ {heading, doan, gach_dau_dong, danh_so, trich_dan, code}. Khối `code` giữ nguyên văn
    trong `chu` và KHÔNG bị diễn giải inline — code là vùng bảo vệ, sửa khéo trong đó là sửa sai.
    """
    dong_list = _bo_frontmatter(van_ban.replace("\r\n", "\n").split("\n"))
    khoi = []
    dem_doan = []
    trong_code = False
    dem_code = []

    def chot_doan():
        if dem_doan:
            khoi.append({"loai": "doan", "muc": 0, "doan": chia_inline(" ".join(dem_doan))})
            dem_doan.clear()

    for dong in dong_list:
        if RAO_CODE.match(dong):
            if trong_code:
                khoi.append({"loai": "code", "muc": 0, "chu": "\n".join(dem_code)})
                dem_code = []
                trong_code = False
            else:
                chot_doan()
                trong_code = True
            continue
        if trong_code:
            dem_code.append(dong)
            continue

        if not dong.strip():
            chot_doan()
            continue
        if DUONG_KE.match(dong):
            chot_doan()
            continue

        khop = HEADING.match(dong)
        if khop:
            chot_doan()
            muc = min(len(khop.group(1)), 3)
            khoi.append({"loai": "heading", "muc": muc, "doan": chia_inline(khop.group(2).strip())})
            continue

        khop = TRICH_DAN.match(dong)
        if khop:
            chot_doan()
            khoi.append({"loai": "trich_dan", "muc": 0, "doan": chia_inline(khop.group(1).strip())})
            continue

        khop = GACH_DAU_DONG.match(dong)
        if khop:
            chot_doan()
            khoi.append(
                {"loai": "gach_dau_dong", "muc": 0, "doan": chia_inline(khop.group(1).strip())}
            )
            continue

        khop = DANH_SO.match(dong)
        if khop:
            chot_doan()
            khoi.append({"loai": "danh_so", "muc": 0, "doan": chia_inline(khop.group(1).strip())})
            continue

        dem_doan.append(dong.strip())

    if trong_code and dem_code:
        khoi.append({"loai": "code", "muc": 0, "chu": "\n".join(dem_code)})
    chot_doan()
    return khoi


# ---------------------------------------------------------------- dựng docx

THIEU_THU_VIEN = (
    "Thiếu thư viện `python-docx`. Cài rồi chạy lại:\n"
    "    pip install python-docx\n"
    "Script này không có đường vòng: docx là định dạng nhị phân, không ghi tay được."
)


def _nap_docx():
    """Nạp python-docx muộn để parser vẫn import/test được trên máy chưa cài thư viện."""
    try:
        import docx
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.shared import Cm, Pt
    except ImportError as loi:  # pragma: no cover - phụ thuộc môi trường
        raise RuntimeError(THIEU_THU_VIEN) from loi
    return docx, WD_ALIGN_PARAGRAPH, qn, Cm, Pt


def _dat_font(run, qn, Pt, ten_font, co_chu, dam=False, nghieng=False):
    run.font.name = ten_font
    run.font.size = Pt(co_chu)
    run.bold = dam
    run.italic = nghieng
    # Word đọc `w:eastAsia` riêng; không đặt thì dấu tiếng Việt có thể rơi sang font khác.
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), ten_font)


def dung_tai_lieu(khoi, quy_cach=None):
    """Danh sách khối ➜ đối tượng `Document` đã áp quy cách. Không ghi đĩa."""
    quy_cach = {**QUY_CACH, **(quy_cach or {})}
    docx, WD_ALIGN_PARAGRAPH, qn, Cm, Pt = _nap_docx()

    tai_lieu = docx.Document()

    binh_thuong = tai_lieu.styles["Normal"]
    binh_thuong.font.name = quy_cach["font"]
    binh_thuong.font.size = Pt(quy_cach["co_chu"])
    binh_thuong.element.get_or_add_rPr().get_or_add_rFonts().set(
        qn("w:eastAsia"), quy_cach["font"]
    )
    binh_thuong.paragraph_format.line_spacing = quy_cach["gian_dong"]

    for phan in tai_lieu.sections:
        phan.top_margin = Cm(quy_cach["le_tren_cm"])
        phan.bottom_margin = Cm(quy_cach["le_duoi_cm"])
        phan.left_margin = Cm(quy_cach["le_trai_cm"])
        phan.right_margin = Cm(quy_cach["le_phai_cm"])

    for muc in khoi:
        loai = muc["loai"]
        if loai == "code":
            doan_van = tai_lieu.add_paragraph()
            doan_van.paragraph_format.line_spacing = 1.0
            run = doan_van.add_run(muc["chu"])
            _dat_font(run, qn, Pt, quy_cach["font_ma"], quy_cach["co_chu"] - 1)
            continue

        if loai == "heading":
            doan_van = tai_lieu.add_paragraph()
            co_chu = quy_cach["co_heading"][muc["muc"]]
            dam_mac_dinh, nghieng_mac_dinh = True, False
            doan_van.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif loai == "gach_dau_dong":
            doan_van = tai_lieu.add_paragraph(style="List Bullet")
            co_chu, dam_mac_dinh, nghieng_mac_dinh = quy_cach["co_chu"], False, False
        elif loai == "danh_so":
            doan_van = tai_lieu.add_paragraph(style="List Number")
            co_chu, dam_mac_dinh, nghieng_mac_dinh = quy_cach["co_chu"], False, False
        elif loai == "trich_dan":
            doan_van = tai_lieu.add_paragraph()
            doan_van.paragraph_format.left_indent = Cm(1.0)
            co_chu, dam_mac_dinh, nghieng_mac_dinh = quy_cach["co_chu"], False, True
        else:
            doan_van = tai_lieu.add_paragraph()
            doan_van.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            co_chu, dam_mac_dinh, nghieng_mac_dinh = quy_cach["co_chu"], False, False

        doan_van.paragraph_format.line_spacing = quy_cach["gian_dong"]
        for phan_chu in muc["doan"]:
            run = doan_van.add_run(phan_chu.chu)
            _dat_font(
                run,
                qn,
                Pt,
                quy_cach["font_ma"] if phan_chu.ma else quy_cach["font"],
                co_chu,
                dam=phan_chu.dam or dam_mac_dinh,
                nghieng=phan_chu.nghieng or nghieng_mac_dinh,
            )
    return tai_lieu


def duong_dan_dich(nguon, out):
    """`--out` là thư mục hay là file? Thư mục thì lấy tên theo file nguồn."""
    dich = Path(out)
    if dich.suffix.lower() == ".docx":
        return dich
    if dich.suffix:
        raise ValueError(
            f"`--out` phải là thư mục hoặc đường dẫn kết thúc bằng .docx, nhận được: {out}"
        )
    return dich / (Path(nguon).stem + ".docx")


def xuat(nguon, out, provenance=None, quy_cach=None):
    """Chạy trọn một lượt: đọc md ➜ dựng docx ➜ ghi đĩa (+ chép sidecar). Trả về đường dẫn docx."""
    nguon = Path(nguon)
    van_ban = nguon.read_text(encoding="utf-8")
    dich = duong_dan_dich(nguon, out)
    dich.parent.mkdir(parents=True, exist_ok=True)
    dung_tai_lieu(phan_tich_markdown(van_ban), quy_cach).save(str(dich))

    if provenance:
        canh_ben = dich.with_suffix("")
        canh_ben = canh_ben.with_name(canh_ben.name + ".provenance.json")
        shutil.copyfile(str(Path(provenance)), str(canh_ben))
    return dich


def main(argv=None):
    if hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8")
    if hasattr(sys.stderr, "reconfigure"):
        sys.stderr.reconfigure(encoding="utf-8")

    bo_doc = argparse.ArgumentParser(
        description="Markdown ➜ .docx theo quy cách văn bản tiếng Việt."
    )
    bo_doc.add_argument("nguon", help="file .md cần xuất")
    bo_doc.add_argument(
        "--out",
        required=True,
        help="thư mục của người dùng, hoặc đường dẫn .docx cụ thể",
    )
    bo_doc.add_argument(
        "--provenance",
        default=None,
        help="sidecar tự khai nguồn gốc, sẽ được chép sang cạnh file docx",
    )
    tham_so = bo_doc.parse_args(argv)

    nguon = Path(tham_so.nguon)
    if not nguon.is_file():
        print(f"Không thấy file nguồn: {nguon}", file=sys.stderr)
        return EXIT_LOI
    if tham_so.provenance and not Path(tham_so.provenance).is_file():
        print(f"Không thấy sidecar provenance: {tham_so.provenance}", file=sys.stderr)
        return EXIT_LOI

    try:
        dich = xuat(nguon, tham_so.out, tham_so.provenance)
    except (RuntimeError, ValueError) as loi:
        print(str(loi), file=sys.stderr)
        return EXIT_LOI

    print(f"Đã xuất: {dich.resolve()}")
    if tham_so.provenance:
        print(f"Sidecar provenance: {dich.with_name(dich.stem + '.provenance.json').resolve()}")
    else:
        print(
            "⚠️  Không có sidecar provenance đi kèm. Bản giao của studio phải mang theo bản tự khai "
            "nguồn gốc — chạy lại với `--provenance` khi bài đã qua trục 2 hoặc trục 4."
        )
    return EXIT_OK


if __name__ == "__main__":
    sys.exit(main())
