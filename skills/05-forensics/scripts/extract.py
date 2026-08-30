#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
extract.py — Bước 0: trích văn bản + metadata, gán ID/offset cho từng câu.

    python extract.py "bai.docx" --out .work/

Sinh: .work/text.txt · .work/meta.json · .work/sentences.json

Phụ thuộc: python-docx (cho .docx), pymupdf (cho .pdf). Cả hai TÙY CHỌN —
thiếu thì chỉ xử lý được .txt.
"""
import argparse, importlib.util, json, re, sys, zipfile
from pathlib import Path

_VI_SEGMENT_PATH = Path(__file__).with_name("vi_segment.py")
_VI_SEGMENT_SPEC = importlib.util.spec_from_file_location("forensics_vi_segment", _VI_SEGMENT_PATH)
if _VI_SEGMENT_SPEC is None or _VI_SEGMENT_SPEC.loader is None:
    raise RuntimeError(f"Không thể nạp bộ tách câu bắt buộc: {_VI_SEGMENT_PATH}")
_VI_SEGMENT = importlib.util.module_from_spec(_VI_SEGMENT_SPEC)
_VI_SEGMENT_SPEC.loader.exec_module(_VI_SEGMENT)

def from_docx(p):
    import docx
    d = docx.Document(str(p))
    parts = [x.text.strip() for x in d.paragraphs if x.text.strip()]
    for t in d.tables:
        for row in t.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    cp = d.core_properties
    meta = {f: (str(getattr(cp, f)) if getattr(cp, f, None) is not None else None)
            for f in ["author", "last_modified_by", "revision", "created",
                      "modified", "title", "category", "last_printed"]}
    try:
        with zipfile.ZipFile(str(p)) as z:
            app = z.read("docProps/app.xml").decode("utf-8", "ignore")
        for tag in ["TotalTime", "Words", "Pages", "Characters", "Application",
                    "AppVersion", "Company", "Template"]:
            m = re.search(rf"<{tag}>(.*?)</{tag}>", app)
            if m:
                v = m.group(1)
                meta[tag] = int(v) if v.isdigit() else v
        with zipfile.ZipFile(str(p)) as z:
            if "word/settings.xml" in z.namelist():
                s = z.read("word/settings.xml").decode("utf-8", "ignore")
                meta["rsid_distinct"] = len(set(re.findall(r'w:val="([0-9A-Fa-f]{8})"', s)))
    except Exception as e:
        meta["_app_xml_error"] = str(e)
    # PII bên thứ ba: che một phần, không đưa nguyên văn vào báo cáo
    if isinstance(meta.get("Company"), str) and re.search(r"[\w.]+@[\w.]+|\d{9,}", meta["Company"]):
        meta["Company_contains_pii"] = True
        meta["Company"] = re.sub(r"([\w.]{2})[\w.]*(@[\w.]+)", r"\1***\2",
                                 re.sub(r"\d{5,}", "***", meta["Company"]))
    return "\n".join(parts), meta

def from_pdf(p):
    import fitz
    doc = fitz.open(str(p))
    pages = [pg.get_text() for pg in doc]
    text = "\n".join(pages)
    chars = sum(len(t.strip()) for t in pages)
    meta = dict(doc.metadata or {})
    meta["Pages"] = len(pages)
    # Ít chữ trên nhiều trang => nhiều khả năng là bản scan cần OCR
    meta["ocr"] = chars < 200 * len(pages)
    return text, meta

def assert_full_coverage(text, sents):
    """sentences.json phải phủ 100% văn bản.

    Bản trước bỏ im lặng câu dưới 15 ký tự, nên `machine_written_spans[]` của trục 2 có thể hợp lệ
    theo schema mà vẫn không nói gì về hai câu có thật (ca `.work/cot-b-ai-baitap`: 45/47). Phép
    kiểm dưới đây là bất biến, không phải cảnh báo: nối `text` của mọi câu, bỏ khoảng trắng, phải
    ra đúng văn bản gốc đã bỏ khoảng trắng.
    """
    joined = re.sub(r"\s+", "", "".join(s["text"] for s in sents))
    whole = re.sub(r"\s+", "", text)
    if joined != whole:
        raise RuntimeError(
            "Bộ tách câu làm mất văn bản: {} ký tự có nghĩa trong sentences.json so với {} của bản "
            "gốc. Không được ghi sentences.json thiếu câu — ID sẽ trỏ sai ở mọi trục sau."
            .format(len(joined), len(whole))
        )


def sentences(text):
    """Tách câu an toàn với viết tắt học thuật (PGS. TS., tr., v.v., 1.1., 3.14).

    Giữ MỌI câu, kể cả câu ngắn (`short: true`). Câu bị bỏ là câu không có id, và câu không có id
    là câu không ai khai được — xem docs/results/self-audit-cot-B.md §2.1.
    """
    raw = _VI_SEGMENT.split_sentences(text)
    out = [{"id": f"s{i+1:04d}", **s,
            "n_syllables": _VI_SEGMENT.count_orthographic_syllables(s["text"])}
           for i, s in enumerate(raw)]
    assert_full_coverage(text, out)
    return out

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("path")
    ap.add_argument("--out", default=".work")
    a = ap.parse_args()
    p, out = Path(a.path), Path(a.out)
    out.mkdir(parents=True, exist_ok=True)
    ext = p.suffix.lower()
    if ext == ".docx":
        text, meta = from_docx(p)
    elif ext == ".pdf":
        text, meta = from_pdf(p)
    else:
        text, meta = p.read_text(encoding="utf-8"), {}
    meta.setdefault("ocr", False)
    meta["source_file"] = p.name
    sents = sentences(text)
    (out / "text.txt").write_text(text, encoding="utf-8")
    (out / "meta.json").write_text(json.dumps(meta, ensure_ascii=False, indent=1, default=str),
                                   encoding="utf-8")
    (out / "sentences.json").write_text(json.dumps(sents, ensure_ascii=False, indent=1),
                                        encoding="utf-8")
    sys.stdout.reconfigure(encoding="utf-8")
    short = sum(1 for s in sents if s.get("short"))
    print(f"-> {out}/  |  {len(text)} ký tự · {len(sents)} câu ({short} câu ngắn, vẫn có id) "
          f"· ocr={meta['ocr']}")
    if len(re.findall(r'\S+', text)) < 300:
        print("!! CẢNH BÁO: dưới 300 âm tiết -> insufficient_evidence, dừng giám định")

if __name__ == "__main__":
    main()
